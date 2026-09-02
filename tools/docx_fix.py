#!/usr/bin/env python3
"""
docx_fix.py — 기획서의 낡은 숫자·용어를 산출물 기준으로 고친다.

── 왜 별도 도구인가 ────────────────────────────────────────────
`docx_check.py` 는 어긋남을 **찾을** 뿐이다. docx 는 텍스트가 run 단위로
쪼개져 있어 손으로 고치면 반드시 빠뜨린다 — 같은 문장이 표 안에 다시
나오는 자리가 여섯 곳이었다.

★ R9(문자열 치환 패처 금지)의 예외가 아니다. R9 가 막는 것은 **소스
  코드**를 스크립트로 고치는 것이다. docx 는 git diff 가 안 되는 바이너리라
  사람이 diff 로 검토할 수 없고, 그래서 도구가 필요한 반대 경우다.

IN    docs/*.docx · data/golden/segments.fingerprint.json
OUT   docs/*.docx (제자리 수정)
PARAM --write 없이는 아무것도 쓰지 않는다
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
from firelane import paths as _p

GOLDEN = _p.GOLDEN / "segments.fingerprint.json"


def rules() -> list[tuple[str, str, str]]:
    """(정규식, 치환, 사유)."""
    g = json.loads(GOLDEN.read_text(encoding="utf-8"))["L1"]
    v = g["verdict"]
    n = g["n"]
    out = [
        # ★ 2026-09-02. 종전 규칙이 `구간` 만 봤는데 기획서는 같은 것을
        #   `산출단위` 로 적는다. **존재하지 않는 형태를 찾고 있었다** —
        #   잡는 쪽(docx_check)과 고치는 쪽이 같이 비껴갔다(DECISIONS §92).
        #
        # ★ 단위 낱말을 조건으로 걸지 않는다. `1,102` 와 `산출단위` 가 서로
        #   다른 run 에 있으면 run 단위 치환에서 패턴이 성립하지 않는다 —
        #   워드는 편집 이력이나 맞춤법 표시만으로도 한 문장을 쪼갠다.
        #   실제로 셋 중 하나(`P379`)가 그래서 남았다. 이 문서에서 `1,102` 는
        #   구간 수 말고 쓰일 자리가 없으므로 수만 보고 바꾼다.
        (r"1,?102", f"{n:,}", "노드접합·산출단위 병합 + 수치지형도 교체"),
        (r"src/etl/pipeline\.py", "src/firelane/pipeline.py", "패키지화 2026-08-21"),
        (r"src/etl\b", "src/firelane", "패키지화 2026-08-21"),
        (r"requirements-etl\.txt[^/]*/\s*", "", "삭제됨. uv.lock 이 정본"),
    ]
    for label, key in (("통행 불가", "blocked"), ("통행 가능", "clear"),
                       ("판정 보류", "needs_cv"), ("영상판정 불가", "unknown")):
        out.append((rf"({label})\s*\d{{2,4}}", rf"\g<1> {v[key]}", "golden 기준"))

    # ── PostGIS 미채택 (docker-compose.yml 주석이 근거) ──
    # segments.geojson 996KB · 1,101구간이라 PostGIS 를 쓸 규모가 아니고,
    # 파이프라인이 285초에 결정론적으로 재생성되므로 DB 의 주 가치인
    # 상태 보존이 필요 없다. DB 가 필요해지는 시점은 동적 계층뿐이며
    # 그때도 seg_uid → status 한 테이블이면 된다(PLAN §2-2).
    out += [
        (r"PostgreSQL\s*\+\s*PostGIS\s*정적·동적\s*계층\s*분리\s*적재",
         "정적 계층 파일(GeoJSON) + 동적 상태 테이블 분리", "PostGIS 미채택"),
        (r"FastAPI,\s*PostGIS,\s*프론트엔드를",
         "FastAPI, 상태 DB, 프론트엔드를", "PostGIS 미채택"),
        (r"PostGIS\s*적재", "동적 상태 테이블 적재", "PostGIS 미채택"),
        (r"PostgreSQL\s*\+\s*PostGIS", "동적 상태 테이블", "PostGIS 미채택"),
        (r"(?<![\w])PostGIS(?![\w])", "상태 DB", "PostGIS 미채택"),
    ]
    # ── 판정 모델 (PLAN §12-4 · MASTER §2-2 · §3-4) ──
    # ★ 2026-08-31. 기획서가 **폐기된 설계**를 적고 있었다. 실물
    #   `segments.geojson` 에는 `width_m` 도 `tier` 도 없다 —
    #     실물   width_min_m(도로경계 최소) · width_max_m(벽~벽 최대)
    #            verdict = clear / needs_cv / blocked / unknown
    #     기획서 width_m(중앙값) · width_min_m(하위 5% 분위수)
    #            tier = FIXED_PASS / CANDIDATE / FIXED_BLOCK · 임계 12.0/2.0
    #
    #   `docx_check` 가 이것을 못 잡은 이유는 **판정 숫자만 보기 때문**이다.
    #   필드명과 어휘는 대상이 아니었다. 아래 규칙과 함께 RETIRED 에도 넣는다.
    out += [
        (r"하위\s*5%\s*분위수를\s*width_min_m\s*으?로\s*확정한다",
         "도로경계 기준 최소 폭을 width_min_m 으로, 건물 폴리곤 기준 "
         "벽~벽 최대 폭을 width_max_m 으로 확정한다", "MASTER §3-4"),
        (r"남은\s*샘플의\s*중앙값을\s*width_m\(표시용\),\s*하위\s*5%\s*분위수를\s*"
         r"width_min_m\(판정용\)으로\s*확정한다",
         "남은 샘플로 width_min_m(도로경계 기준 최소 폭)을 내고, 건물 폴리곤에서 "
         "width_max_m(벽~벽 최대 폭)을 따로 낸다. 둘의 중앙 차이는 +2.08m 이며 "
         "그 사이에 전신주·화분·실외기·주차 차량이 들어 있다", "MASTER §3-4"),
        # ★ 표 셀은 `<w:t>` 하나에 담긴다. 앞 셀(필드명)까지 이어 잡으려 하면
        #   xml 에서 안 맞는다 — 평문으로만 이어 보이는 것이다.
        (r"중심선\s*샘플\s*폭의\s*중앙값\s*\(표시용\)",
         "건물 폴리곤 기준 벽~벽 최대 폭. 605/1,101 만 값이 있다", "MASTER §3-4"),
        (r"교차로\s*5m\s*제외\s*후\s*하위\s*5%\s*분위수\s*\(판정용\)",
         "도로경계 기준 최소 폭. 확실히 비어 있는 폭", "MASTER §3-4"),
        # 필드명 셀 자체
        (r"^width_m$", "width_max_m", "MASTER §3-4"),
        (r"width_min_m\s*12\.0m\s*이상은\s*FIXED_PASS,\s*2\.0m\s*미만은\s*"
         r"FIXED_BLOCK,\s*그\s*사이는\s*CANDIDATE로\s*분류하여\s*CANDIDATE만\s*"
         r"영상\s*판정\s*대상으로\s*한다",
         "최대 폭이 3.0m 미만이면 blocked, 최소 폭이 7.0m 이상이면 clear, "
         "그 사이는 needs_cv 로 분류하여 needs_cv 만 영상 판정 대상으로 한다. "
         "7.0 = 통과 하한 3.0 + 양쪽 주차 2 x 2.0 이다", "MASTER §2-2"),
        (r"width_m\s*/\s*width_min_m", "width_min_m / width_max_m", "MASTER §3-4"),
        (r"세그먼트\s*도형,\s*width_m,\s*width_min_m,",
         "세그먼트 도형, width_min_m, width_max_m,", "MASTER §3-4"),
    ]

    # ── tier 표 (셀 단위) ──
    #   실물 verdict 는 clear / needs_cv / blocked / unknown 이고
    #   tier 어휘는 존재하지 않는다. 임계도 7.0 / 3.0 이다(MASTER §2-2).
    out += [
        (r"≥\s*12\.0\s*m", "최소 폭 ≥ 7.0 m", "MASTER §2-2"),
        (r"2\.0\s*~\s*12\.0\s*m", "그 사이", "MASTER §2-2"),
        (r"&lt;\s*2\.0\s*m", "최대 폭 &lt; 3.0 m", "MASTER §2-2"),
        (r"3단\s*tier\s*분류\s*\(FIXED_PASS\s*/\s*CANDIDATE\s*/\s*FIXED_BLOCK\)",
         "판정 4종 (clear / needs_cv / blocked / unknown)", "MASTER §2-2"),
        (r"3단\s*tier\s*분류", "판정 4종 분류", "MASTER §2-2"),
        (r"FIXED_PASS", "clear", "MASTER §2-2"),
        (r"FIXED_BLOCK", "blocked", "MASTER §2-2"),
        (r"CANDIDATE", "needs_cv", "MASTER §2-2"),
        (r"tiertextFIXED_PASS\s*\|\s*CANDIDAT", "verdicttextclear | needs_cv", "MASTER §2-2"),
        # "12.0m 상한은 팀 자체 추정치" — 임계 자체가 없어졌다.
        (r"12\.0m\s*상한\s*임계값은\s*팀\s*자체\s*추정치이며\s*발표\s*시\s*그\s*사실을\s*명시한다",
         "통과 하한 3.0m 는 소방청 「2025 화재현장 골든타임 확보 종합대책」의 "
         "진입불가 기준(폭 2m 이하)에 차량 전폭 2.5m + 미러·조향 여유 0.5m 를 "
         "더한 값이다. clear 하한 7.0m 는 여기에 양쪽 주차 2 x 2.0m 를 더한 것이다",
         "MASTER §2-2"),
        (r"상한\s*12\.0m는\s*팀\s*자체\s*추정치\(4\.0\s*→\s*8\.0\s*→\s*12\.0으로\s*조정\)이며,\s*"
         r"발표\s*시\s*자체\s*추정임을\s*명시한다",
         "clear 하한 7.0m 는 통과 하한 3.0m 에 양쪽 주차 2 x 2.0m 를 더한 값이며, "
         "양방 주차가 상시인 골목에서 한쪽만 가정하면 판정이 낙관 쪽으로 기운다",
         "MASTER §2-2"),
    ]

    # ── 2차 정리 (셀 원문 확인 후) ──
    # ★ 1차에서 27 run 을 고쳤으나 여섯 자리가 남았다. `&lt; 2.0 m` 은 셀
    #   하나에 통째로 들어 있었고 내 정규식이 공백을 잘못 잡았다.
    #   **평문으로 보이는 것과 xml 실물이 다르다** — 셀 원문을 뜬 뒤에야 맞았다.
    out += [
        (r"&lt;\s*2\.0\s*m", "최대 폭 &lt; 3.0 m", "MASTER §2-2"),
        (r"하한\s*2\.0m는\s*소방청\s*진입불가\s*기준과\s*일치하므로\s*외부\s*방어가\s*가능하다",
         "blocked 하한 3.0m 는 소방청 진입불가 기준(폭 2m 이하)에 차량 전폭 2.5m + "
         "미러·조향 여유 0.5m 를 더한 값이라 외부 방어가 가능하다", "MASTER §2-2"),
        (r"needs_cv그\s*사이영상\s*판정\s*대상",
         "needs_cv최소 폭 3.0 ~ 7.0 m영상 판정 대상", "MASTER §2-2"),
        # 폭 산출 절차 — 분위수는 폐기됐다
        (r"수치지도\s*트랜섹트\s*하위\s*5%\s*분위수를\s*판정\s*입력으로\s*사용",
         "수치지도 트랜섹트 도로경계 기준 최소 폭을 판정 입력으로 사용", "MASTER §3-4"),
        (r"폭\s*측정에는\s*최댓값이\s*아니라\s*하위\s*5%\s*분위수를\s*사용하는데",
         "폭은 최소(도로경계)와 최대(벽~벽) 둘을 함께 내는데", "MASTER §3-4"),
        (r"교차로\s*반경\s*5m\s*샘플\s*제외\s*후\s*하위\s*5%\s*분위수\s*산출",
         "교차로 반경 5m 샘플 제외 후 최소·최대 폭 산출", "MASTER §3-4"),
        (r"교차로\s*제외\s*분위수\s*산출", "교차로 제외 최소·최대 폭 산출", "MASTER §3-4"),
        (r"교차로\s*제외\s*분위수노드", "교차로 제외 폭 산출노드", "MASTER §3-4"),
        (r"최댓값이\s*아니라\s*하위\s*분위수를\s*쓰는\s*이유는\s*소방차가\s*구간의\s*평균\s*폭이\s*"
         r"아니라\s*가장\s*좁은\s*병목을\s*통과하지\s*못하기\s*때문이다",
         "판정에 최소 폭을 쓰는 이유는 소방차가 구간의 평균 폭이 아니라 "
         "가장 좁은 병목을 통과하지 못하기 때문이다", "MASTER §3-4"),
        # tier → verdict (실물 필드명)
        (r"tiertextclear\s*\|\s*needs_cv\s*\|\s*blocked",
         "verdicttextclear | needs_cv | blocked | unknown", "실물 필드"),
        (r"width_src,\s*tier,\s*cctv_coverage", "width_src, verdict, cctv_coverage", "실물 필드"),
        (r"폭\s*소스,\s*tier,\s*현재\s*상태", "폭 소스, verdict, 현재 상태", "실물 필드"),
        (r"폭,\s*폭\s*소스,\s*tier,\s*CCTV", "폭, 폭 소스, verdict, CCTV", "실물 필드"),
        (r"bbox\s*및\s*tier\s*필터", "bbox 및 verdict 필터", "실물 필드"),
        (r"bbox·tier\s*필터", "bbox·verdict 필터", "실물 필드"),
        (r"\(도형·폭·tier·커버리지\)", "(도형·폭·verdict·커버리지)", "실물 필드"),
        (r"tier\s*값\(clear\)은\s*유지해", "verdict 값(clear)은 유지해", "실물 필드"),
        (r"tier\s*재분류\s*및\s*4색\s*분포\s*확정", "verdict 재분류 및 4색 분포 확정", "실물 필드"),
    ]

    # ── 3차 (셀 경계를 넘지 않는다) ──
    # ★ 2차에서 넷이 또 안 잡혔다. `needs_cv그 사이…` 처럼 **앞 셀까지 이어**
    #   잡으려 했기 때문이다. 평문에서는 이어 보이지만 xml 에서는 각각
    #   다른 `<w:t>` 다. 셀 하나 안의 문자열만으로 좁힌다.
    #   같은 실수를 세 번 했다 — 평문을 보고 정규식을 쓰면 반드시 이렇게 된다.
    out += [
        (r"^그 사이$", "최소 폭 3.0 ~ 7.0 m", "MASTER §2-2"),
        (r"^&lt; 2\.0 m$", "최대 폭 &lt; 3.0 m", "MASTER §2-2"),
        (r"^Tier$", "판정", "실물 필드 — tier 어휘는 없다"),
        (r"교차로 제외 분위수노드", "교차로 제외 폭 산출노드", "MASTER §3-4"),
        (r"tiertextclear \| needs_cv \| blocked",
         "verdicttextclear | needs_cv | blocked | unknown", "실물 필드"),
    ]

    # ── 소방용수 (PLAN §12-7 · MASTER §3-12) ──
    # 기획서는 소화전과 소방용수시설을 한 숫자로 뭉쳐 588 이라 적었다.
    # 둘은 다른 집계다 — 소화전 589 는 소방용수시설 654 의 부분집합이다.
    # 공개율 논거는 소방용수시설 전체가 분모여야 맞다(MASTER §16-3).
    out += [
        (r"소방용수\s*588개\s*\(\s*지상식\s*431\s*[+·]\s*지하식\s*157\s*\)",
         "소화전 589개(지상식 418 · 지하식 171) · 소방용수시설 총계 654개",
         "MASTER §3-12"),
        (r"588개\s*\(\s*지상식\s*431\s*[·+]\s*지하식\s*157\s*\)",
         "654개(소화전 589 = 지상식 418 · 지하식 171)", "MASTER §3-12"),
        (r"588개\s*중\s*31개", "654개 중 31개", "분모는 소방용수시설 총계"),
        (r"5%\s*\(\s*588개\s*중\s*31개\s*\)", "5%(654개 중 31개)", "MASTER §3-12"),
        (r"(?<![\d,])588개", "654개", "MASTER §3-12"),
    ]

    # ── 차량 제원 (PLAN §12-2 · MASTER §3-13) ──
    # 규격은 확보됐다. 남은 추정은 축거·최소회전반경 둘뿐이며 그것이
    # `wheelbase_verified: false` 다. "제원 미확보" 로 뭉뚱그리면 이미 한
    # 근거를 안 한 것으로 적는 셈이고, 외부 독자가 그대로 읽는다.
    _SPEC = ("소방청 「소방장비 기본규격」 소방펌프차 KFS-1-0073-2025-00 §3.3"
             "(2025-12-24 고시)")
    out += [
        (r"정확한 표준 제원은 소방청 소방차량 제작규격 및 소방서 문의로 "
         r"확보 예정이며,\s*그 이전까지는 잠정값을 사용한다\.",
         f"표준 제원은 {_SPEC}으로 확보했다. 중형 기준 전폭 2.5m 이하이며 "
         "통과 하한 3.0m 는 여기에 미러·조향 여유 0.5m 를 더한 값이다. "
         "다만 축거와 최소회전반경은 공식 규격에 없어 추정값으로 남으며, "
         "내륜차 계산에 그 둘이 필요하다.", "MASTER §3-13"),
        (r"현재 세 값 모두 잠정이며 확정 경로가 각각 다르다\.",
         "전폭은 KFS-1-0073-2025-00 §3.3 으로 확정했고, 축거·최소회전반경은 "
         "공식 규격에 없어 추정으로 남는다. 확정 경로가 각각 다르다.",
         "MASTER §3-13"),
        (r"차량 제원 미확보", "축거·회전반경 미확보", "규격은 확보됨"),
        (r"판정 통과선의 근거가 잠정값에 머무름",
         "내륜차 계산의 근거가 추정값에 머무름", "MASTER §3-13"),
        (r"소방청 소방차량 제작규격 확보 및 동부소방서 현장대응과 문의 병행",
         f"{_SPEC} 확보 완료. 축거·회전반경은 동부소방서 인터뷰로 확정",
         "MASTER §3-13"),
        (r"제원 확보 시 재검증한다\.",
         "축거·회전반경 확보 시 내륜차를 재산출한다.", "MASTER §3-13"),
    ]

    return out


def fix(p: Path, write: bool) -> int:
    import docx
    d = docx.Document(str(p))
    R = rules()
    n = 0

    def do(par) -> int:
        """run 단위로 먼저, 안 되면 문단 단위로.

        ★ docx 는 텍스트가 run 으로 임의 분할된다. `PostGIS 적재` 가
          `PostG` + `IS 적재` 두 run 에 걸쳐 있으면 run 단위 치환이
          조용히 통과한다 — 이 저장소가 계속 겪은 조용한 실패다.
          그래서 문단 전체로 한 번 더 본다.
        """
        c = 0
        for run in par.runs:
            t0 = run.text
            t = t0
            for rx, rep, _ in R:
                t = re.sub(rx, rep, t)
            if t != t0:
                run.text = t
                c += 1

        # run 경계를 넘는 구절
        whole = par.text
        fixed = whole
        for rx, rep, _ in R:
            fixed = re.sub(rx, rep, fixed)
        if fixed != whole and par.runs:
            # ★ 첫 run 에 몰아넣으면 그 문단의 부분 서식이 사라진다.
            #   숫자·용어 교정이 서식보다 중요하므로 감수하되 세어서 알린다.
            par.runs[0].text = fixed
            for r in par.runs[1:]:
                r.text = ""
            c += 1
        return c

    for par in d.paragraphs:
        n += do(par)
    for t in d.tables:
        for r in t.rows:
            for cell in r.cells:
                for par in cell.paragraphs:
                    n += do(par)

    if write and n:
        d.save(str(p))
    return n


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--write", action="store_true")
    a = ap.parse_args()
    total = 0
    for p in sorted((ROOT / "docs").glob("*.docx")):
        c = fix(p, a.write)
        total += c
        print(f"  {p.name}  run {c}개 {'수정' if a.write else '수정 예정'}")
    if not a.write:
        print("\n아무것도 안 썼다. 적용하려면 --write")
        print("★ docx 는 git diff 가 안 된다. 적용 후 워드로 눈으로 확인할 것.")
    else:
        print(f"\n{total}개 run 수정. tools/docx_check.py 로 대조해라.")
    return 0


if __name__ == "__main__":
    sys.exit(main())

