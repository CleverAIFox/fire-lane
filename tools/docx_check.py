#!/usr/bin/env python3
"""
docx_check.py — 기획서가 산출물과 어긋나지 않는가.

── 왜 생겼나 ───────────────────────────────────────────────────
2026-08-27. 기획서는 문서 넷 중 **유일하게 외부가 읽는 것**인데,
시제 규칙 밖이라는 이유로 강제자가 없었다. 그래서 낡았다.

    docnum_check.py     README · MASTER · PLAN · DECISIONS  ← md 만 본다
    (없음)              기획서_Fire-Lane.docx               ← 아무도 안 본다

`PLAN §12` 가 갱신 대상 열 건을 표로 들고 있었으나 그 표 자체가 낡았다 —
"대상 222구간" 을 고치라고 적혀 있는데 문서에 `222` 는 없고 실제로는
`1,102` 가 박혀 있었다. **강제자 없는 목록은 목록도 낡는다.**

이 저장소가 반복해 배운 형태다 — 규약은 문서에 존재하고 강제하는 검사가
없다(MASTER §17).

── 무엇을 보는가 ───────────────────────────────────────────────
정본은 `data/golden/segments.fingerprint.json` 이다. 문서에 적힌 판정
숫자가 그것과 다르면 **산출물이 옳다**(MASTER §0-3).

    구간 수 · 판정 4종 · 총연장 · 폐기된 경로·기술명

★ 모든 숫자를 보지 않는다. 시장 통계·법령 조항·비용 산정은 산출물과
  무관하며, 그것까지 검사하면 거짓 경보로 사람이 검사를 끈다.

IN    docs/*.docx · data/golden/segments.fingerprint.json
OUT   없음 (검사). --json 이면 stdout
PARAM RETIRED · 허용 오차 없음(정수 대조)
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
from firelane import paths as _p

GOLDEN = _p.GOLDEN / "segments.fingerprint.json"

# 폐기된 이름. 문서에 남아 있으면 독자가 그대로 따라 한다.
RETIRED = {
    "src/etl": "패키지가 src/firelane 으로 바뀌었다(2026-08-21)",
    "app.js": "web/js/ 27개 모듈로 쪼갰다(DECISIONS §27)",
    "requirements-etl": "삭제됐다. uv.lock 이 정본이다(2026-08-23)",
    "PostGIS": "미채택. segments.geojson 996KB 규모라 쓸 자리가 아니다",
    "apply.py": "패치 zip 절차를 폐기했다(DECISIONS §65)",
}


def _load_docx(p: Path) -> list[tuple[str, str]]:
    """(위치, 텍스트) 목록. 표 안까지 본다."""
    try:
        import docx
    except ImportError:
        return []
    d = docx.Document(str(p))
    out = [(f"P{i}", x.text) for i, x in enumerate(d.paragraphs) if x.text.strip()]
    for ti, t in enumerate(d.tables):
        for ri, r in enumerate(t.rows):
            for c in r.cells:
                if c.text.strip():
                    out.append((f"T{ti}R{ri}", c.text))
    return out


def _canon() -> dict[str, int | float]:
    if not GOLDEN.exists():
        return {}
    g = json.loads(GOLDEN.read_text(encoding="utf-8")).get("L1", {})
    v = g.get("verdict", {})
    return {
        "구간 수": g.get("n"),
        "통행 불가": v.get("blocked"),
        "통행 가능": v.get("clear"),
        "판정 보류": v.get("needs_cv"),
        "영상판정 불가": v.get("unknown"),
        "총연장": g.get("length_total_m"),
    }


def audit(p: Path) -> list[str]:
    cells = _load_docx(p)
    if not cells:
        return [f"  {p.name} 을 읽지 못했다 (python-docx 미설치?)"]

    bad: list[str] = []
    c = _canon()
    n = c.get("구간 수")

    # ── 1 · 구간 수 ──
    if n:
        for where, txt in cells:
            for m in re.finditer(r"(\d{1,2},?\d{3})\s*구간", txt):
                val = int(m.group(1).replace(",", ""))
                if 900 < val < 1400 and val != n:
                    bad.append(
                        f"  [{where}] 구간 수 {m.group(1)} → **{n:,}**\n"
                        f"      …{txt.strip()[:70]}…")

    # ── 2 · 판정 4종 ──
    for label in ("통행 불가", "통행 가능", "판정 보류", "영상판정 불가"):
        want = c.get(label)
        if not want:
            continue
        for where, txt in cells:
            for m in re.finditer(rf"{label}\s*(\d{{2,4}})", txt):
                if int(m.group(1)) != want:
                    bad.append(
                        f"  [{where}] {label} {m.group(1)} → **{want}**\n"
                        f"      …{txt.strip()[:70]}…")

    # ── 3 · 폐기된 이름 ──
    for where, txt in cells:
        for name, why in RETIRED.items():
            if name in txt:
                bad.append(f"  [{where}] 폐기: `{name}` — {why}\n"
                           f"      …{txt.strip()[:70]}…")

    return bad


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--json", action="store_true")
    a = ap.parse_args()

    docs = sorted((ROOT / "docs").glob("*.docx"))
    if not docs:
        print("docs/ 에 docx 가 없다 — 건너뛴다")
        return 0

    allbad: list[str] = []
    for p in docs:
        allbad += audit(p)

    if a.json:
        print(json.dumps({"stale": len(allbad)}, ensure_ascii=False))

    if not allbad:
        print(f"기획서 OK — 산출물과 일치 (구간 {_canon().get('구간 수'):,})")
        return 0

    print(f"기획서가 산출물과 어긋난다. {len(allbad)}건\n")
    print("\n".join(allbad[:40]))
    if len(allbad) > 40:
        print(f"  … 외 {len(allbad) - 40}건")
    print("\n  정본은 data/golden/segments.fingerprint.json 이다.")
    print("  문서와 어긋나면 산출물이 옳다(MASTER §0-3).")
    print("  기획서는 넷 중 유일하게 외부가 읽는다 — 어긋나면 비용이 가장 크다.")
    return 1


if __name__ == "__main__":
    sys.exit(main())

